"""
Extracts plain text from uploaded profile/resume files.
Supports: PDF, DOCX, TXT.
"""
import io
from fastapi import HTTPException

ALLOWED_EXTENSIONS = {"pdf", "docx", "txt"}


def get_extension(filename: str) -> str:
    if "." not in filename:
        return ""
    return filename.rsplit(".", 1)[-1].lower()


def validate_extension(filename: str) -> str:
    ext = get_extension(filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=415,
            detail=f"Unsupported file type '.{ext}'. Allowed types: PDF, DOCX, TXT.",
        )
    return ext


def extract_text(file_bytes: bytes, ext: str) -> str:
    try:
        if ext == "pdf":
            return _extract_pdf(file_bytes)
        if ext == "docx":
            return _extract_docx(file_bytes)
        if ext == "txt":
            return file_bytes.decode("utf-8", errors="ignore")
    except HTTPException:
        raise
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(
            status_code=422,
            detail=f"Could not read the uploaded file. It may be corrupted or password protected. ({exc})",
        ) from exc
    raise HTTPException(status_code=415, detail="Unsupported file type.")


def _extract_pdf(file_bytes: bytes) -> str:
    import pdfplumber

    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    return "\n".join(text_parts)


def _extract_docx(file_bytes: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(file_bytes))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)
